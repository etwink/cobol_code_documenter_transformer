"""
Build document clusters for hierarchical summarization.

Strategy:
  1. COBOL files: grouped by dependency graph (entry-point programs + their
     transitive dependencies form one cluster).
  2. Word/Excel files: first checked for mentions of COBOL program names and
     merged into the matching COBOL cluster; remaining docs are clustered by
     an LLM call that groups them into logical business domains.
"""

import json
import re
from dataclasses import dataclass, field
from pathlib import Path

from cobol_dependency_analyzer import (
    parse_file,
    build_dependency_graph,
    find_transitive_dependencies,
)

# Soft trigger: when a cluster reaches this many files, run a context-window
# check to determine the true safe limit before splitting.
_MAX_CLUSTER_FILES = 15

# Never split a cluster below this floor (keeps sub-clusters meaningful).
_MIN_CLUSTER_FILES = 5

# Hard ceiling regardless of how small the individual files are.
_MAX_CLUSTER_FILES_ABSOLUTE = 50

# Total character budget for ONE synthesis LLM call (input side).
# ~400k chars ≈ 85k tokens — leaves ample room for reasoning + output tokens
# on a 200k-token context model.
_CONTEXT_WINDOW_LIMIT_CHARS = 400_000

# Fixed overhead of the synthesis prompt template (header, instructions, dep
# info, word-count directive).  Measured from the actual HierarchicalSummarizer
# synthesis prompts; add a generous buffer for cluster/entry-point names.
_SYNTHESIS_TEMPLATE_CHARS = 2_000

# Chars read per file when estimating cluster content size (mirrors the
# excerpt size used in _summarize_files_individually).
_FILE_EXCERPT_CHARS = 6_000

# Multi-cluster doc assignment: a doc is added to any cluster whose hit count is
# at least this fraction of the best (highest-scoring) cluster's hit count.
_MULTI_CLUSTER_RATIO = 0.5

# Global doc threshold: if a doc qualifies for this many or more clusters it is
# treated as a shared reference document rather than a member of any one cluster.
# It will be injected as context into every matching cluster's synthesis prompt.
_GLOBAL_CLUSTER_THRESHOLD = 3


@dataclass
class DocumentCluster:
    cluster_id: str
    cluster_name: str
    # "cobol", "docs", or "mixed" (COBOL cluster that absorbed related Word/Excel docs)
    cluster_type: str
    cobol_files: list[Path] = field(default_factory=list)
    doc_files: list[Path] = field(default_factory=list)
    # Stem names of COBOL programs in this cluster (upper-case), used for cross-referencing
    cobol_program_names: set[str] = field(default_factory=set)
    # Entry-point program name for COBOL clusters (None for pure-doc clusters)
    entry_point: str | None = None
    # Docs that matched multiple clusters and are injected as shared context
    # rather than assigned as members.  Not counted in file_count.
    shared_doc_files: list[Path] = field(default_factory=list)

    @property
    def all_files(self) -> list[Path]:
        return self.cobol_files + self.doc_files

    @property
    def file_count(self) -> int:
        return len(self.all_files)


class ClusterBuilder:
    """Build document clusters from a scanned document set."""

    def __init__(self):
        from llm_integration import AzureLLMClient
        self.llm = AzureLLMClient()

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def build_clusters(
        self,
        cobol_files: list[Path],
        word_files: list[Path],
        excel_files: list[Path],
        code_files: list[Path] | None = None,
        context_block: str = "",
    ) -> list[DocumentCluster]:
        """
        Return a list of DocumentClusters ready for hierarchical summarization.

        Steps:
          1. Group COBOL files by dependency graph.
          2. Match Word/Excel/code docs to COBOL clusters by program-name mentions.
          3. LLM-cluster the remaining unmatched docs.
        """
        clusters = self._build_cobol_clusters(cobol_files, context_block=context_block)
        all_docs = (code_files or []) + word_files + excel_files
        unmatched_docs = self._attach_docs_to_cobol_clusters(clusters, all_docs)
        if unmatched_docs:
            doc_clusters = self._llm_cluster_docs(unmatched_docs, context_block=context_block)
            clusters.extend(doc_clusters)
        return clusters

    # ------------------------------------------------------------------
    # COBOL clustering via dependency graph
    # ------------------------------------------------------------------

    def _build_cobol_clusters(self, cobol_files: list[Path], context_block: str = "") -> list[DocumentCluster]:
        if not cobol_files:
            return []

        analyses = [parse_file(p) for p in cobol_files]
        graph = build_dependency_graph(analyses)

        # Identify entry points: files that appear as a source but are never
        # someone else's target (i.e. nothing calls them).
        all_targets: set[str] = {
            d["target"] for deps in graph.values() for d in deps
        }
        all_sources: set[str] = set(graph.keys())
        entry_points = all_sources - all_targets

        # Files that don't appear in the graph at all are standalone.
        graphed_names = set(graph.keys()) | all_targets
        standalone_names = {
            p.stem.upper() for p in cobol_files
        } - graphed_names

        # Build path lookup: stem (upper) → Path
        path_by_stem: dict[str, Path] = {
            p.stem.upper(): p for p in cobol_files
        }

        clusters: list[DocumentCluster] = []
        assigned: set[str] = set()

        for ep in sorted(entry_points):
            transitive = find_transitive_dependencies(graph, ep)
            cluster_stems = {ep} | transitive

            cluster_files = [
                path_by_stem[s] for s in cluster_stems if s in path_by_stem
            ]
            if not cluster_files:
                continue

            clusters.append(
                DocumentCluster(
                    cluster_id=f"cobol_{ep.lower()}",
                    cluster_name=f"COBOL Subsystem: {ep}",
                    cluster_type="cobol",
                    cobol_files=cluster_files,
                    cobol_program_names=cluster_stems,
                    entry_point=ep,
                )
            )
            assigned |= cluster_stems

        # Standalone / unassigned COBOL files (pure copybooks, utilities)
        unassigned_paths = [
            p for p in cobol_files
            if p.stem.upper() not in assigned
        ]
        # Also add standalone programs not connected to anything
        standalone_paths = [
            path_by_stem[s] for s in standalone_names if s in path_by_stem
        ]
        leftover = list({p for p in unassigned_paths + standalone_paths})
        if leftover:
            stems = {p.stem.upper() for p in leftover}
            clusters.append(
                DocumentCluster(
                    cluster_id="cobol_utilities",
                    cluster_name="COBOL Shared Utilities & Copybooks",
                    cluster_type="cobol",
                    cobol_files=sorted(leftover),
                    cobol_program_names=stems,
                    entry_point=None,
                )
            )

        # Split any cluster that exceeds the soft trigger.  The context-window
        # check computes how many files actually fit given the synthesis prompt
        # template overhead + the user-provided context block length.
        result: list[DocumentCluster] = []
        for cl in clusters:
            if cl.file_count > _MAX_CLUSTER_FILES:
                max_size = _compute_cluster_file_limit(cl.all_files, context_block)
                result.extend(self._split_large_cluster(cl, graph, path_by_stem, max_size=max_size))
            else:
                result.append(cl)
        return result

    def _split_large_cluster(
        self,
        cluster: DocumentCluster,
        graph: dict,
        path_by_stem: dict[str, Path],
        _depth: int = 0,
        max_size: int = _MAX_CLUSTER_FILES,
    ) -> list[DocumentCluster]:
        """
        Recursively split an oversized COBOL cluster into sub-clusters.

        For clusters WITH an entry point: use the entry point's direct callees
        as sub-cluster roots; each gets its transitive deps (bounded to the
        parent file set).  The entry point + leftover form a residual core.

        For clusters WITHOUT an entry point (utilities bucket): find internal
        "mini-entry-points" by building a sub-graph restricted to these files,
        then fall back to alphabetical chunking if no structure is found.

        Recursion is capped at depth 3 to guarantee termination.
        """
        if cluster.file_count <= max_size or _depth >= 3:
            return [cluster]

        parent_stems = cluster.cobol_program_names

        if cluster.entry_point:
            # ── Has entry point: split by direct callees ──────────────────────
            ep = cluster.entry_point
            direct_deps = sorted(
                {d["target"] for d in graph.get(ep, [])} & parent_stems
            )
            roots = direct_deps
        else:
            # ── No entry point: find internal structure ───────────────────────
            # Build a sub-graph restricted to this cluster's files
            sub_graph: dict[str, list] = {
                k: [d for d in v if d["target"] in parent_stems]
                for k, v in graph.items()
                if k in parent_stems
            }
            sub_targets = {d["target"] for deps in sub_graph.values() for d in deps}
            sub_sources = set(sub_graph.keys())
            # Mini-entry-points: present in sub-graph as callers but never as callees
            mini_entries = sub_sources - sub_targets
            # Also include stems that have no graph presence at all — treat
            # them as independent units so they still get split into chunks.
            if not mini_entries:
                # Fully disconnected or cyclic — fall back to alphabetical chunks
                return self._split_into_chunks(cluster, path_by_stem, max_size=max_size)
            roots = sorted(mini_entries)

        # ── Build sub-clusters from roots ─────────────────────────────────────
        sub_clusters: list[DocumentCluster] = []
        absorbed: set[str] = set()

        for root in roots:
            if root in absorbed:
                continue
            sub_stems = ({root} | find_transitive_dependencies(graph, root)) & parent_stems
            sub_files = [path_by_stem[s] for s in sorted(sub_stems) if s in path_by_stem]
            if not sub_files:
                continue
            sub_clusters.append(
                DocumentCluster(
                    cluster_id=f"cobol_{root.lower()}",
                    cluster_name=f"COBOL Subsystem: {root}",
                    cluster_type="cobol",
                    cobol_files=sub_files,
                    cobol_program_names=sub_stems,
                    entry_point=root,
                )
            )
            absorbed |= sub_stems

        # Residual: entry point itself + anything not absorbed
        residual_stems = parent_stems - absorbed
        if residual_stems:
            residual_files = [path_by_stem[s] for s in sorted(residual_stems) if s in path_by_stem]
            if residual_files:
                ep = cluster.entry_point
                sub_clusters.append(
                    DocumentCluster(
                        cluster_id=f"cobol_{ep.lower()}_core" if ep else "cobol_utilities_core",
                        cluster_name=f"COBOL Core: {ep}" if ep else "COBOL Shared Utilities",
                        cluster_type="cobol",
                        cobol_files=residual_files,
                        cobol_program_names=residual_stems,
                        entry_point=ep,
                    )
                )

        if not sub_clusters:
            return self._split_into_chunks(cluster, path_by_stem, max_size=max_size)

        # Recurse on any sub-cluster that is still too large
        result: list[DocumentCluster] = []
        for sc in sub_clusters:
            if sc.file_count > max_size:
                result.extend(self._split_large_cluster(sc, graph, path_by_stem, _depth + 1, max_size=max_size))
            else:
                result.append(sc)
        return result

    def _split_into_chunks(
        self,
        cluster: DocumentCluster,
        path_by_stem: dict[str, Path],
        max_size: int = _MAX_CLUSTER_FILES,
    ) -> list[DocumentCluster]:
        """Last-resort: split alphabetically into chunks of max_size."""
        files = sorted(cluster.cobol_files, key=lambda p: p.name)
        chunks: list[DocumentCluster] = []
        for i in range(0, len(files), max_size):
            batch = files[i : i + max_size]
            part = i // max_size + 1
            stems = {p.stem.upper() for p in batch}
            chunks.append(
                DocumentCluster(
                    cluster_id=f"cobol_util_part_{part}",
                    cluster_name=f"COBOL Utilities (Part {part})",
                    cluster_type="cobol",
                    cobol_files=batch,
                    cobol_program_names=stems,
                    entry_point=None,
                )
            )
        return chunks

    # ------------------------------------------------------------------
    # Attach Word/Excel docs to COBOL clusters by name-mention
    # ------------------------------------------------------------------

    def _attach_docs_to_cobol_clusters(
        self,
        cobol_clusters: list[DocumentCluster],
        doc_files: list[Path],
    ) -> list[Path]:
        """
        For each doc file, score it against every COBOL cluster by counting
        program-name mentions.  Three outcomes are possible:

        1. No matches → unmatched (goes to LLM clustering later).
        2. Matches 1–(_GLOBAL_CLUSTER_THRESHOLD-1) clusters above the ratio
           threshold → added to each matching cluster's doc_files (multi-cluster
           assignment).
        3. Matches _GLOBAL_CLUSTER_THRESHOLD or more clusters → treated as a
           global/shared reference document and added to each matching cluster's
           shared_doc_files instead, so it informs all those summaries without
           inflating cluster membership counts.

        Returns the list of doc files that didn't match any cluster.
        """
        if not doc_files or not cobol_clusters:
            return doc_files

        unmatched: list[Path] = []

        for doc_path in doc_files:
            text = _read_text_excerpt(doc_path, max_chars=50_000)
            if not text:
                unmatched.append(doc_path)
                continue

            # Score every cluster
            scored: list[tuple[int, DocumentCluster]] = []
            for cluster in cobol_clusters:
                hits = _count_program_mentions(text, cluster.cobol_program_names)
                if hits > 0:
                    scored.append((hits, cluster))

            if not scored:
                unmatched.append(doc_path)
                continue

            best_hits = max(h for h, _ in scored)
            min_hits = max(1, int(best_hits * _MULTI_CLUSTER_RATIO))
            matching = [c for h, c in scored if h >= min_hits]

            if len(matching) >= _GLOBAL_CLUSTER_THRESHOLD:
                # Global doc: inject as shared context, not a cluster member
                for cluster in matching:
                    cluster.shared_doc_files.append(doc_path)
            else:
                # Assign to each qualifying cluster
                for cluster in matching:
                    cluster.doc_files.append(doc_path)
                    if cluster.cluster_type == "cobol":
                        cluster.cluster_type = "mixed"

        return unmatched

    # ------------------------------------------------------------------
    # LLM clustering for remaining docs
    # ------------------------------------------------------------------

    def _llm_cluster_docs(self, doc_files: list[Path], context_block: str = "") -> list[DocumentCluster]:
        """Ask the LLM to group unmatched docs into logical business domains."""
        if not doc_files:
            return []

        # Build a short description list for the LLM
        entries: list[str] = []
        for p in doc_files:
            excerpt = _read_text_excerpt(p, max_chars=400).replace("\n", " ").strip()
            entries.append(f'- "{p.name}": {excerpt[:300]}')

        prompt = _build_doc_clustering_prompt(entries, context_block=context_block)
        try:
            response = self.llm.query(prompt, temperature=0, max_tokens=6000)
            groups = _parse_clustering_response(response, doc_files)
        except Exception:
            groups = []

        # If LLM produced only one group (or failed entirely), fall back to a
        # file-type-based split so we at least get some differentiation.
        if len(groups) <= 1:
            groups = _fallback_cluster_by_type(doc_files)

        clusters: list[DocumentCluster] = []
        for i, group in enumerate(groups):
            clusters.append(
                DocumentCluster(
                    cluster_id=f"docs_{i}",
                    cluster_name=group["name"],
                    cluster_type="docs",
                    doc_files=group["files"],
                )
            )
        return clusters


# ------------------------------------------------------------------
# Helpers
# ------------------------------------------------------------------

def _compute_cluster_file_limit(files: list[Path], context_block: str = "") -> int:
    """
    Compute how many files can fit in one synthesis LLM call without overflowing
    the context window.

    Budget = _CONTEXT_WINDOW_LIMIT_CHARS
           - _SYNTHESIS_TEMPLATE_CHARS   (fixed prompt overhead)
           - len(context_block)           (user-provided context injected into prompt)
           ÷ avg_file_excerpt_chars       (sampled from the actual files)

    Result is clamped to [_MIN_CLUSTER_FILES, _MAX_CLUSTER_FILES_ABSOLUTE].
    """
    template_chars = _SYNTHESIS_TEMPLATE_CHARS + len(context_block)
    available_chars = _CONTEXT_WINDOW_LIMIT_CHARS - template_chars
    if available_chars <= 0:
        return _MIN_CLUSTER_FILES

    # Sample up to 5 files for a better average estimate
    sample = files[:5]
    total_chars = sum(
        len(_read_text_excerpt(p, max_chars=_FILE_EXCERPT_CHARS)) or _FILE_EXCERPT_CHARS
        for p in sample
    )
    avg_chars = total_chars // max(len(sample), 1)
    limit = available_chars // max(avg_chars, 100)
    return max(_MIN_CLUSTER_FILES, min(_MAX_CLUSTER_FILES_ABSOLUTE, limit))


def _read_text_excerpt(path: Path, max_chars: int = 50_000) -> str:
    """Return a plain-text excerpt from a file (best-effort)."""
    try:
        suffix = path.suffix.upper()
        if suffix == ".DOCX":
            from docx import Document as DocxDocument
            doc = DocxDocument(path)
            text = "\n".join(p.text for p in doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    text += " ".join(c.text for c in row.cells) + "\n"
            return text[:max_chars]
        elif suffix in {".XLSX", ".XLS", ".XLSM"}:
            from openpyxl import load_workbook
            wb = load_workbook(path, read_only=True, data_only=True)
            parts: list[str] = []
            for ws in wb.worksheets:
                parts.append(f"Sheet: {ws.title}")
                for row in ws.iter_rows(values_only=True):
                    parts.append(
                        " | ".join(str(c) for c in row if c is not None)
                    )
            return "\n".join(parts)[:max_chars]
        elif suffix == ".XLSB":
            try:
                from pyxlsb import open_workbook  # type: ignore
                parts: list[str] = []
                with open_workbook(str(path)) as wb:
                    for sheet_name in wb.sheets:
                        parts.append(f"Sheet: {sheet_name}")
                        with wb.get_sheet(sheet_name) as ws:
                            for row in ws.rows():
                                cells = [str(c.v) for c in row if c.v is not None]
                                if cells:
                                    parts.append(" | ".join(cells))
                return "\n".join(parts)[:max_chars]
            except ImportError:
                return ""
        else:
            return path.read_text(encoding="utf-8", errors="ignore")[:max_chars]
    except Exception:
        return ""


def _count_program_mentions(text: str, program_names: set[str]) -> int:
    """Count how many distinct program names from the set appear in text."""
    text_upper = text.upper()
    return sum(1 for name in program_names if re.search(r"\b" + re.escape(name) + r"\b", text_upper))


def _build_doc_clustering_prompt(file_entries: list[str], context_block: str = "") -> str:
    n = len(file_entries)
    # Aim for roughly one cluster per 6–7 files, capped at 6 so the result stays
    # manageable.  The LLM is allowed to use fewer if documents don't naturally
    # divide into that many topics.
    target_groups = max(2, min(6, n // 6))
    entries_text = "\n".join(file_entries)
    context_section = f"\nProcess context provided by the user:\n{context_block}\n" if context_block else ""
    return f"""You are organizing {n} files (business documents and/or source code) into logical groups for analysis.
{context_section}
Below are the filenames and short excerpts. Group them into approximately {target_groups}–{target_groups + 2} clusters based on shared business domain, process area, or subject matter. Use fewer clusters if the documents clearly belong to fewer topics — do not force artificial splits.

Files:
{entries_text}

Return ONLY valid JSON in this exact format (no markdown, no explanation):
{{
  "clusters": [
    {{"name": "Descriptive Group Name", "files": ["filename1.docx", "filename2.xlsx"]}},
    ...
  ]
}}"""


def _fallback_cluster_by_type(doc_files: list[Path]) -> list[dict]:
    """
    Last-resort fallback when LLM clustering fails or returns only one group.
    Splits files into buckets by broad file type so there is at least some
    differentiation rather than one giant cluster.
    """
    word_exts = {".DOCX", ".DOC"}
    excel_exts = {".XLSX", ".XLS", ".XLSM", ".XLSB"}
    code_exts = {".PY", ".SQL", ".JS", ".TS", ".VB", ".BAS", ".CS", ".JAVA",
                 ".PS1", ".R", ".SH", ".BAT", ".CMD"}

    buckets: dict[str, list[Path]] = {}
    for p in doc_files:
        ext = p.suffix.upper()
        if ext in word_exts:
            buckets.setdefault("Business Documents (Word)", []).append(p)
        elif ext in excel_exts:
            buckets.setdefault("Spreadsheets & Workbooks (Excel)", []).append(p)
        elif ext in code_exts:
            buckets.setdefault("Source Code & Scripts", []).append(p)
        else:
            buckets.setdefault("Supporting Documents", []).append(p)

    return [{"name": name, "files": files} for name, files in buckets.items()]


def _parse_clustering_response(
    response: str,
    all_files: list[Path],
) -> list[dict]:
    """Parse LLM JSON response into a list of {{name, files}} dicts."""
    path_by_name: dict[str, Path] = {p.name: p for p in all_files}

    # Strip markdown fences if present
    clean = re.sub(r"```(?:json)?|```", "", response).strip()
    try:
        data = json.loads(clean)
        groups: list[dict] = []
        assigned: set[str] = set()

        for cluster in data.get("clusters", []):
            name = cluster.get("name", "Documents")
            files: list[Path] = []
            for fname in cluster.get("files", []):
                if fname in path_by_name and fname not in assigned:
                    files.append(path_by_name[fname])
                    assigned.add(fname)
            if files:
                groups.append({"name": name, "files": files})

        # Any files not placed by the LLM go into a catch-all
        leftover = [p for p in all_files if p.name not in assigned]
        if leftover:
            groups.append({"name": "General Documents", "files": leftover})

        return groups
    except (json.JSONDecodeError, KeyError):
        return [{"name": "Business Documents", "files": all_files}]


