"""Document loaders for various file formats."""

from pathlib import Path
from html.parser import HTMLParser
from docx import Document as DocxDocument
from openpyxl import load_workbook

from .base import BaseDocumentLoader, DocumentContent


class TextDocumentLoader(BaseDocumentLoader):
    """Loader for plain text files."""

    def load(self) -> DocumentContent:
        self._validate_file_size()
        with open(self.file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        return DocumentContent(
            filename=self.file_path.name,
            file_type="text",
            content=content,
            metadata={"encoding": "utf-8"}
        )


class COBOLDocumentLoader(BaseDocumentLoader):
    """Loader for COBOL source files (.CIC, .CPY, .MPS, .SRC, .CT1, .JCV, .PRV, .COB, .CBL)."""

    def load(self) -> DocumentContent:
        self._validate_file_size()
        with open(self.file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        return DocumentContent(
            filename=self.file_path.name,
            file_type="cobol",
            content=content,
            metadata={
                "extension": self.file_path.suffix,
                "is_fixed_format": self._is_fixed_format(content)
            }
        )

    @staticmethod
    def _is_fixed_format(content: str) -> bool:
        """Detect if COBOL file is in fixed or free format."""
        lines = content.split('\n')[:20]  # Check first 20 lines
        fixed_format_count = 0
        for line in lines:
            if len(line) > 7 and line[6] in ['*', '/', ' ', 'D']:
                fixed_format_count += 1
        return fixed_format_count > len(lines) * 0.5


class WordDocumentLoader(BaseDocumentLoader):
    """Loader for Microsoft Word documents (.docx)."""

    def load(self) -> DocumentContent:
        self._validate_file_size()
        doc = DocxDocument(self.file_path)
        content = "\n".join([para.text for para in doc.paragraphs])
        
        # Include table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                content += f"\n{row_text}"
        
        return DocumentContent(
            filename=self.file_path.name,
            file_type="docx",
            content=content,
            metadata={"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)}
        )


class ExcelDocumentLoader(BaseDocumentLoader):
    """Loader for Excel files — reads both computed cell values AND formula expressions."""

    # Cap how many formula cells we capture per sheet to avoid very large prompts.
    _MAX_FORMULA_CELLS_PER_SHEET = 150

    def load(self) -> DocumentContent:
        self._validate_file_size()

        # ── Pass 1: computed values (data_only=True) ──────────────────────────
        sheet_names: list[str] = []
        values_sections: list[str] = []
        try:
            wb_values = load_workbook(self.file_path, read_only=True, data_only=True)
            sheet_names = list(wb_values.sheetnames)
            for sheet_name in sheet_names:
                sheet = wb_values[sheet_name]
                rows: list[str] = []
                for row in sheet.iter_rows(values_only=True):
                    if any(c is not None for c in row):
                        rows.append(" | ".join(str(c) if c is not None else "" for c in row))
                values_sections.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows))
            wb_values.close()
        except Exception:
            pass

        # ── Pass 2: formula expressions (data_only=False) ────────────────────
        formula_sections: list[str] = []
        has_formulas = False
        try:
            wb_formula = load_workbook(self.file_path, data_only=False)
            for sheet_name in wb_formula.sheetnames:
                sheet = wb_formula[sheet_name]
                formula_lines: list[str] = []
                count = 0
                for row in sheet.iter_rows():
                    for cell in row:
                        if isinstance(cell.value, str) and cell.value.startswith("="):
                            formula_lines.append(f"  {cell.coordinate}: {cell.value}")
                            has_formulas = True
                            count += 1
                            if count >= self._MAX_FORMULA_CELLS_PER_SHEET:
                                formula_lines.append(
                                    f"  ... (capped at {self._MAX_FORMULA_CELLS_PER_SHEET} formulas for this sheet)"
                                )
                                break
                    if count >= self._MAX_FORMULA_CELLS_PER_SHEET:
                        break
                if formula_lines:
                    formula_sections.append(
                        f"=== Sheet: {sheet_name} ===\n" + "\n".join(formula_lines)
                    )
            wb_formula.close()
        except Exception:
            pass

        values_text = "\n\n".join(values_sections) if values_sections else "(no data)"
        if has_formulas:
            formulas_text = "\n\n".join(formula_sections)
            content = (
                "--- DATA VALUES (computed results as seen in Excel) ---\n"
                f"{values_text}\n\n"
                "--- FORMULAS (calculation logic that produces those results) ---\n"
                f"{formulas_text}"
            )
        else:
            content = f"--- DATA VALUES ---\n{values_text}"

        return DocumentContent(
            filename=self.file_path.name,
            file_type="excel",
            content=content,
            metadata={"sheets": sheet_names, "has_formulas": has_formulas},
        )


class MacroEnabledExcelDocumentLoader(ExcelDocumentLoader):
    """Loader for macro-enabled Excel workbooks (.xlsm).

    Extends ExcelDocumentLoader with a third pass that extracts VBA macro
    source code using oletools.olevba.  Falls back gracefully if oletools
    is not installed or the file contains no macros.
    """

    # Cap total VBA output so it doesn't overwhelm the context window.
    _MAX_VBA_CHARS = 20_000

    def load(self) -> DocumentContent:
        # Pass 1 + 2: sheet values and formulas (inherited)
        base = super().load()

        # Pass 3: VBA macro source code
        vba_section = self._extract_vba()

        content = base.content
        if vba_section:
            content += f"\n\n--- VBA MACROS (automation logic embedded in the workbook) ---\n{vba_section}"

        return DocumentContent(
            filename=base.filename,
            file_type="excel_macro",
            content=content,
            metadata={**base.metadata, "has_vba": bool(vba_section)},
        )

    def _extract_vba(self) -> str:
        try:
            from oletools.olevba import VBA_Parser  # type: ignore
        except ImportError:
            return "(oletools not installed — install it with: pip install oletools)"

        try:
            vba_parser = VBA_Parser(str(self.file_path))
            if not vba_parser.detect_vba_macros():
                return ""

            parts: list[str] = []
            total_chars = 0
            for (_filename, _stream_path, vba_filename, vba_code) in vba_parser.extract_macros():
                if not vba_code or not vba_code.strip():
                    continue
                header = f"=== Module: {vba_filename} ==="
                chunk = f"{header}\n{vba_code.strip()}"
                if total_chars + len(chunk) > self._MAX_VBA_CHARS:
                    remaining = self._MAX_VBA_CHARS - total_chars
                    if remaining > len(header) + 20:
                        parts.append(chunk[:remaining] + "\n... (truncated)")
                    break
                parts.append(chunk)
                total_chars += len(chunk)
            vba_parser.close()
            return "\n\n".join(parts)
        except Exception:
            return ""


class BinaryExcelDocumentLoader(BaseDocumentLoader):
    """Loader for Excel Binary Workbooks (.xlsb).

    .xlsb files use a proprietary binary format that openpyxl cannot read.
    - Data values are extracted via pyxlsb.
    - Formula expressions are NOT available (stored as compiled binary).
    - VBA macros are extracted via oletools.olevba (same as .xlsm).
    """

    _MAX_VBA_CHARS = 20_000

    def load(self) -> DocumentContent:
        self._validate_file_size()

        values_text = self._extract_values()
        vba_text = self._extract_vba()

        parts = ["--- DATA VALUES (cell values as seen in Excel) ---\n" + values_text]
        parts.append(
            "--- FORMULAS ---\n"
            "(Formula expressions are not available for binary .xlsb workbooks — "
            "only computed values can be read.)"
        )
        if vba_text:
            parts.append(
                "--- VBA MACROS (automation logic embedded in the workbook) ---\n" + vba_text
            )

        return DocumentContent(
            filename=self.file_path.name,
            file_type="excel_binary",
            content="\n\n".join(parts),
            metadata={"has_vba": bool(vba_text)},
        )

    def _extract_values(self) -> str:
        try:
            from pyxlsb import open_workbook  # type: ignore
        except ImportError:
            return "(pyxlsb not installed — install it with: pip install pyxlsb)"

        try:
            sections: list[str] = []
            with open_workbook(str(self.file_path)) as wb:
                for sheet_name in wb.sheets:
                    with wb.get_sheet(sheet_name) as sheet:
                        rows: list[str] = []
                        for row in sheet.rows():
                            cells = [str(c.v) if c.v is not None else "" for c in row]
                            if any(cells):
                                rows.append(" | ".join(cells))
                        sections.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows))
            return "\n\n".join(sections) if sections else "(no data)"
        except Exception as e:
            return f"(error reading values: {e})"

    def _extract_vba(self) -> str:
        try:
            from oletools.olevba import VBA_Parser  # type: ignore
        except ImportError:
            return ""

        try:
            vba_parser = VBA_Parser(str(self.file_path))
            if not vba_parser.detect_vba_macros():
                return ""

            parts: list[str] = []
            total_chars = 0
            for (_filename, _stream_path, vba_filename, vba_code) in vba_parser.extract_macros():
                if not vba_code or not vba_code.strip():
                    continue
                header = f"=== Module: {vba_filename} ==="
                chunk = f"{header}\n{vba_code.strip()}"
                if total_chars + len(chunk) > self._MAX_VBA_CHARS:
                    remaining = self._MAX_VBA_CHARS - total_chars
                    if remaining > len(header) + 20:
                        parts.append(chunk[:remaining] + "\n... (truncated)")
                    break
                parts.append(chunk)
                total_chars += len(chunk)
            vba_parser.close()
            return "\n\n".join(parts)
        except Exception:
            return ""


class HTMLDocumentLoader(BaseDocumentLoader):
    """Loader for HTML files."""

    def load(self) -> DocumentContent:
        self._validate_file_size()
        with open(self.file_path, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = f.read()
        
        # Parse HTML and extract text
        parser = SimpleHTMLParser()
        parser.feed(html_content)
        content = parser.get_text()
        
        return DocumentContent(
            filename=self.file_path.name,
            file_type="html",
            content=content,
            metadata={"raw_html_length": len(html_content)}
        )


class SimpleHTMLParser(HTMLParser):
    """Simple HTML parser to extract text content."""

    def __init__(self):
        super().__init__()
        self.text_content = []
        self.skip_content = False

    def handle_starttag(self, tag, attrs):
        if tag in ['script', 'style']:
            self.skip_content = True

    def handle_endtag(self, tag):
        if tag in ['script', 'style']:
            self.skip_content = False
        elif tag in ['p', 'div', 'br', 'li']:
            self.text_content.append('\n')

    def handle_data(self, data):
        if not self.skip_content:
            text = data.strip()
            if text:
                self.text_content.append(text)

    def get_text(self):
        return ' '.join(self.text_content)


def get_loader(file_path: Path) -> BaseDocumentLoader:
    """Factory function to get appropriate loader for file type."""
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    loaders = {
        # Documents
        '.docx': WordDocumentLoader,
        '.doc': WordDocumentLoader,
        '.xlsx': ExcelDocumentLoader,
        '.xls': ExcelDocumentLoader,
        '.xlsm': MacroEnabledExcelDocumentLoader,
        '.xlsb': BinaryExcelDocumentLoader,
        '.html': HTMLDocumentLoader,
        '.htm': HTMLDocumentLoader,
        # COBOL source
        '.cobol': COBOLDocumentLoader,
        '.cbl': COBOLDocumentLoader,
        '.cob': COBOLDocumentLoader,
        '.cic': COBOLDocumentLoader,
        '.cpy': COBOLDocumentLoader,
        '.mps': COBOLDocumentLoader,
        '.src': COBOLDocumentLoader,
        '.ct1': COBOLDocumentLoader,
        '.jcv': COBOLDocumentLoader,
        '.prv': COBOLDocumentLoader,
        # Other source code — all read as plain text
        '.py': TextDocumentLoader,
        '.sql': TextDocumentLoader,
        '.js': TextDocumentLoader,
        '.ts': TextDocumentLoader,
        '.vb': TextDocumentLoader,
        '.bas': TextDocumentLoader,
        '.cs': TextDocumentLoader,
        '.java': TextDocumentLoader,
        '.ps1': TextDocumentLoader,
        '.r': TextDocumentLoader,
        '.sh': TextDocumentLoader,
        '.bat': TextDocumentLoader,
        '.cmd': TextDocumentLoader,
        '.txt': TextDocumentLoader,
    }

    loader_class = loaders.get(suffix, TextDocumentLoader)
    return loader_class(file_path)
